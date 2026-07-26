from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "vehicle_report_assets"
OUTPUT = OUT_DIR / "Vehicle_Control_Code_and_Strategy_Report.docx"
ARCH_FIG = ASSET_DIR / "vehicle_control_architecture.png"

NAVY = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF3EC"
PALE_GOLD = "FFF4D6"
GRAY = "5B6570"
WHITE = "FFFFFF"
BLACK = "000000"
RED = "9B1C1C"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_para_spacing(paragraph, before=0, after=6, line=1.10, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph):
    paragraph.add_run("Page ")
    add_field(paragraph, "PAGE")
    paragraph.add_run(" of ")
    add_field(paragraph, "NUMPAGES")


def add_paragraph(doc, text="", *, bold_lead=None, italic=False, color=BLACK, after=6, align=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=after)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=color)
        set_run(p.add_run(text[len(bold_lead):]), italic=italic, color=color)
    else:
        set_run(p.add_run(text), italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_para_spacing(p, after=4, line=1.167)
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_spacing(p, after=4, line=1.167)
    set_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=8, line=1.10)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    set_run(p.add_run(f"{label}: "), bold=True, color=NAVY)
    set_run(p.add_run(text), color=NAVY)
    return p


def add_code_block(doc, lines):
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        set_para_spacing(p, before=2 if idx == 0 else 0, after=2 if idx < len(lines)-1 else 8, line=1.0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F7F8FA")
        p_pr.append(shd)
        set_run(p.add_run(line), size=9.2, color=NAVY, font="Consolas")


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0, line=1.0)
        set_run(p.add_run(header), size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            set_para_spacing(p, after=0, line=1.05)
            set_run(p.add_run(str(value)), size=9.3)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def create_architecture_figure():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    scale = 2
    image = Image.new("RGB", (2300, 880), "white")
    draw = ImageDraw.Draw(image)
    font_dir = Path("C:/Windows/Fonts")
    title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 28)
    box_title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 22)
    body_font = ImageFont.truetype(str(font_dir / "arial.ttf"), 18)
    boxes = [
        (60, 210, 390, 235, "Sensors", "Ultrasonic distance\nCamera detections"),
        (590, 210, 410, 235, "Interpretation", "Median filtering\nDirection / risk"),
        (1130, 210, 430, 235, "Safety Controller", "Finite-state loop\nStop-first policy"),
        (1740, 210, 410, 235, "Motor Adapter", "Forward / yaw / stop\nHiWonder SDK"),
        (590, 570, 410, 210, "Fallback", "Alternate left/right\nwhen vision is absent"),
        (1130, 570, 430, 210, "Configuration", "Thresholds, timing,\nvalidation bounds"),
        (1740, 570, 410, 210, "Telemetry", "State events and\nRunResult summary"),
    ]
    for x, y, w, h, title, body in boxes:
        face = "#E8EEF5" if y < 500 else "#F2F4F7"
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=face, outline="#2E74B5", width=3)
        draw.text((x + 24, y + 24), title, font=box_title_font, fill="#18324A")
        draw.multiline_text((x + 24, y + 72), body, font=body_font, fill="#273746", spacing=10)

    def arrow(start, end, color="#2E74B5", width=5):
        draw.line((start, end), fill=color, width=width)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            points = [(ex, ey), (ex - 18, ey - 11), (ex - 18, ey + 11)] if ex > sx else [(ex, ey), (ex + 18, ey - 11), (ex + 18, ey + 11)]
        else:
            points = [(ex, ey), (ex - 11, ey - 18), (ex + 11, ey - 18)] if ey > sy else [(ex, ey), (ex - 11, ey + 18), (ex + 11, ey + 18)]
        draw.polygon(points, fill=color)

    arrow((450, 327), (590, 327))
    arrow((1000, 327), (1130, 327))
    arrow((1560, 327), (1740, 327))
    arrow((795, 570), (1240, 445), color="#5B6570", width=4)
    arrow((1345, 570), (1345, 445), color="#5B6570", width=4)
    arrow((1450, 445), (1945, 570), color="#5B6570", width=4)
    draw.text((60, 55), "Vehicle-side control architecture", font=title_font, fill="#18324A")
    image.save(ARCH_FIG, quality=95)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, after=0, line=1.0)
    set_run(p.add_run("SMART VEHICLE PROJECT  |  VEHICLE CONTROL REPORT"), size=8.5, bold=True, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p, after=0, line=1.0)
    add_page_number(p)
    for run in p.runs:
        set_run(run, size=8.5, color=GRAY)


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_architecture_figure()
    doc = Document()
    configure_document(doc)

    # Editorial technical-report cover.
    p = doc.add_paragraph()
    set_para_spacing(p, before=58, after=16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("TECHNICAL IMPLEMENTATION REPORT"), size=10.5, bold=True, color=BLUE)
    p = doc.add_paragraph()
    set_para_spacing(p, after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Vehicle Control Code and Strategy"), size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_para_spacing(p, after=22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("TurboPi obstacle avoidance, safety control, and visual-guidance integration"), size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    set_para_spacing(p, before=38, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Smart Vehicle Project"), size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_para_spacing(p, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Prepared from the current repository implementation"), size=10.5, italic=True, color=GRAY)
    p = doc.add_paragraph()
    set_para_spacing(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("July 21, 2026"), size=10.5, color=GRAY)

    doc.add_page_break()
    add_heading(doc, "Executive Summary", 1)
    add_paragraph(doc, "The vehicle implementation uses a conservative, layered control strategy. A Raspberry Pi reads the front ultrasonic sensor, filters short-term noise, and drives the TurboPi mecanum chassis through a finite-state controller. The controller always stops before a turn, waits for the chassis to settle, performs a short timed yaw, checks that the path is clear, and only then resumes forward motion. Camera-based guidance is deliberately advisory: it can recommend a turn direction, but it cannot override the ultrasonic stop condition.")
    add_callout(doc, "Core design decision", "Keep the stopping authority in the simple, local ultrasonic loop; use computer vision only to improve the choice of escape direction.", fill=PALE_GREEN)
    add_paragraph(doc, "The present milestone is a bounded 10-second obstacle-avoidance demonstration. It defaults to simulation and requires two explicit command-line flags before real motors are enabled. Four unit tests verify normal completion, final stopping, alternating fallback turns, invalid-sensor fault handling, and side-aware camera guidance. Live-camera logs also show that the visual pipeline can run near real time, although distance calibration and physical-robot trials are still outstanding.")

    add_heading(doc, "1. Scope and System Boundary", 1)
    add_paragraph(doc, "This report covers the vehicle-side code in TurboPiIntegratedGuide and the perception strategy developed in VisualGuideProject. The web backend and Flutter application are outside the real-time motor safety loop. This separation is intentional: network or user-interface failures should not be able to prevent a local emergency stop.")
    add_table(doc,
        ["Layer", "Current role", "Primary modules"],
        [
            ("Hardware interface", "Translate abstract motor and distance operations to the HiWonder SDK.", "hardware.py"),
            ("Safety controller", "Run the timed finite-state avoidance loop and guarantee a final stop.", "controller.py"),
            ("Configuration", "Centralize thresholds, timing, validation bounds, and yaw direction.", "config.py"),
            ("Execution entry point", "Select simulation or hardware, require motion confirmation, and print results.", "main.py"),
            ("Vision adapter", "Convert a selected detection center into an advisory turn direction.", "visual_guidance.py"),
            ("Perception prototype", "Detect, rank, stabilize, and describe obstacles and approach risk.", "VisualGuideProject/*.py"),
        ], [1700, 4560, 3100])

    add_heading(doc, "2. Vehicle-Side Architecture", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=3)
    run = p.add_run()
    picture = run.add_picture(str(ARCH_FIG), width=Inches(6.35))
    picture._inline.docPr.set(
        "descr",
        "Vehicle control architecture from sensors through interpretation and safety control to the motor adapter, with fallback, configuration, and telemetry support.",
    )
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap, after=8, line=1.0)
    set_run(cap.add_run("Figure 1. Local sensing-to-actuation architecture and supporting control inputs."), size=9, italic=True, color=GRAY)
    add_paragraph(doc, "The controller depends on small interfaces rather than concrete hardware classes. Motor and DistanceSensor protocols let the same controller run with real TurboPi adapters or deterministic simulation objects. This makes safety logic testable on a laptop and reduces the amount of code that changes when the hardware implementation changes.")

    add_heading(doc, "3. Obstacle-Avoidance Control Strategy", 1)
    add_heading(doc, "3.1 Finite-state control", 2)
    add_paragraph(doc, "AvoidanceController represents the vehicle behavior as explicit states: STARTING, FORWARD, STOPPING, TURN_LEFT, TURN_RIGHT, CHECK_CLEAR, FINISHED, and FAULT. The state model prevents direct jumps from forward travel into a turn and makes the transition sequence observable through state events.")
    add_code_block(doc, [
        "STARTING -> FORWARD",
        "FORWARD --distance <= obstacle threshold--> STOPPING",
        "STOPPING --settled--> TURN_LEFT or TURN_RIGHT",
        "TURN_* --timer expired--> CHECK_CLEAR",
        "CHECK_CLEAR --distance >= clear threshold--> FORWARD",
        "any exception / repeated invalid data -> FAULT -> stop",
    ])
    add_heading(doc, "3.2 Stop-first maneuver", 2)
    add_paragraph(doc, "When the filtered distance reaches the obstacle threshold, the code increments the obstacle counter and commands an immediate stop. It then waits for stop_settle_seconds before starting the yaw. At the end of the timed turn, it stops again, waits again, and checks the sensor. If the path is still blocked, the controller performs another turn rather than moving forward into uncertainty.")
    add_heading(doc, "3.3 Hysteresis and noise filtering", 2)
    add_paragraph(doc, "Two distance thresholds create hysteresis: the default obstacle threshold is 35 cm, while forward travel resumes only at 45 cm. This 10 cm gap prevents rapid stop/start oscillation near a single threshold. Each valid measurement is added to a three-sample deque and the median is used, reducing the effect of one short ultrasonic spike.")
    add_heading(doc, "3.4 Deterministic turn choice", 2)
    add_paragraph(doc, "If the visual advisor returns left or right, the controller uses that recommendation. If vision is missing, invalid, centered, or raises an exception, the controller alternates left and right. The fallback is deterministic, avoids repeated bias to one side, and does not make vehicle safety dependent on camera availability.")

    add_heading(doc, "4. Safety and Fault-Handling Strategy", 1)
    add_table(doc,
        ["Safeguard", "Implementation", "Purpose"],
        [
            ("Simulation by default", "No SDK commands unless --hardware is supplied.", "Prevents accidental movement during development."),
            ("Two-step motion consent", "Hardware mode also requires --confirm-motion.", "Makes motor enablement deliberate."),
            ("Initial and final stop", "motor.stop() runs before the loop and in finally.", "Establishes a safe initial condition and stops on every exit path."),
            ("Invalid-data fail-safe", "Invalid readings stop the motor; three consecutive failures raise a fault by default.", "Avoids motion when distance cannot be trusted."),
            ("Bounded run", "A monotonic deadline limits the milestone to 10 seconds by default.", "Contains risk during early testing."),
            ("Input validation", "Configuration rejects unsafe or inconsistent values.", "Catches bad tuning before motion begins."),
            ("Yaw reversal option", "--reverse-yaw corrects chassis direction without changing control logic.", "Handles hardware orientation differences safely."),
        ], [2050, 4100, 3210])
    add_callout(doc, "Safety limitation", "The software safety pattern is sound for a prototype, but there is no evidence in the repository of completed physical TurboPi obstacle trials. An accessible physical emergency stop and supervised low-speed testing remain mandatory.", fill=PALE_GOLD)

    add_heading(doc, "5. Visual-Guidance Strategy", 1)
    add_heading(doc, "5.1 Simplified integration boundary", 2)
    add_paragraph(doc, "The integrated vehicle package deliberately reuses only one visual decision: turn away from the selected detection. The frame is divided into thirds. A detection on the left advises a right turn, a detection on the right advises a left turn, and a centered or unavailable detection returns no advice. This narrow interface prevents experimental perception features from entering the motor safety loop.")
    add_heading(doc, "5.2 Perception pipeline used to select and describe hazards", 2)
    add_paragraph(doc, "The broader visual prototype combines fast background/motion candidates with optional YOLO recognition. The fast detector retains warning authority, while YOLO can label a matching candidate. The selected detection is mapped to a 3 x 3 image region, given a rough closeness grade from bounding-box area, and optionally assigned an approximate monocular distance.")
    add_bullet(doc, "Motion and background subtraction provide frequent, lightweight obstacle candidates.")
    add_bullet(doc, "Candidate ranking favors larger and more central objects that are more relevant to the travel path.")
    add_bullet(doc, "YOLO recognition is sampled less often and cached between inference runs to protect frame rate.")
    add_bullet(doc, "Detection and warning stabilizers reduce box jitter and repeated changes in spoken or beep output.")
    add_heading(doc, "5.3 Approach-risk grading", 2)
    add_paragraph(doc, "ApproachRiskTracker uses trends rather than a single frame. Its score considers bounding-box growth, estimated-distance decrease, image-plane speed, side-to-center movement, downward movement, and near-distance grade. Lateral motion without size growth or decreasing distance is capped at low risk, reducing false alarms from objects that merely cross the camera view.")
    add_callout(doc, "Interpretation rule", "Monocular distance and image growth are supporting signals, not physical safety measurements. The ultrasonic sensor remains authoritative for stopping.")

    add_heading(doc, "6. Key Configuration and Tuning Parameters", 1)
    add_table(doc,
        ["Parameter", "Default", "Engineering meaning"],
        [
            ("run_seconds", "10.0 s", "Bounded duration of the early hardware demonstration."),
            ("loop_interval_seconds", "0.05 s", "Nominal 20 Hz controller evaluation interval."),
            ("forward_speed", "25", "Conservative chassis command for initial tests."),
            ("obstacle_distance_cm", "35 cm", "Distance at which forward travel stops."),
            ("clear_distance_cm", "45 cm", "Distance required before forward travel resumes."),
            ("turn_duration_seconds", "0.5 s", "Length of one open-loop yaw maneuver."),
            ("turn_yaw_rate", "0.4", "Magnitude of the TurboPi yaw command."),
            ("stop_settle_seconds", "0.08 s", "Pause before and after turning."),
            ("distance_filter_size", "3 samples", "Median-filter window."),
            ("valid distance range", "2-500 cm", "Rejects implausible sonar values."),
            ("max consecutive invalid", "3", "Fault threshold for unavailable sonar data."),
        ], [2600, 1700, 5060])
    add_paragraph(doc, "The parameter values are deliberately conservative starting points, not final calibrated values. Turn duration and yaw rate are open-loop and must be tuned on the actual floor surface, battery level, load, and chassis orientation.")

    add_heading(doc, "7. Verification Evidence", 1)
    add_heading(doc, "7.1 Automated vehicle-controller tests", 2)
    add_paragraph(doc, "The TurboPiIntegratedGuide test suite was executed during report preparation using the bundled Python runtime. All four tests passed.")
    add_table(doc,
        ["Test", "Verified behavior", "Result"],
        [
            ("Normal completion", "Forward command occurs; final command is stop; state is FINISHED.", "Pass"),
            ("Two obstacles", "Both left and right turns occur through alternating fallback; final command is stop.", "Pass"),
            ("Invalid sensor", "Controller enters FAULT, does not move forward, and finishes with stop.", "Pass"),
            ("Visual advisor", "Left-side detection advises right; right-side advises left; center returns no advice.", "Pass"),
        ], [2100, 5960, 1300])
    add_heading(doc, "7.2 Recorded live-camera evidence", 2)
    add_paragraph(doc, "A 100.1-second visual-guide run recorded on July 13, 2026 processed 2,490 frames at approximately 24.9 FPS. YOLO was active, with 249 inference runs and an average inference time of 112.9 ms. The stabilizer reduced raw warning changes from 136 to 26, an approximately 81% reduction. The same log notes that motion detection was the strongest signal and that distance estimation still requires calibration.")
    add_callout(doc, "Evidence boundary", "The camera log validates perception throughput and stabilization behavior; it does not validate collision avoidance on the physical TurboPi chassis.", fill=PALE_GOLD)

    add_heading(doc, "8. Strengths of the Current Strategy", 1)
    add_bullet(doc, "Safety-critical stopping is local, simple, and independent of the network, backend, and mobile app.")
    add_bullet(doc, "Dependency inversion allows the same controller logic to use simulated or real hardware.")
    add_bullet(doc, "Explicit states make behavior auditable and reduce unsafe transition shortcuts.")
    add_bullet(doc, "Hysteresis, median filtering, and invalid-reading limits address common ultrasonic failure modes.")
    add_bullet(doc, "Vision adds directional intelligence without becoming a single point of failure.")
    add_bullet(doc, "Configuration values are centralized and validated, supporting repeatable tuning.")

    add_heading(doc, "9. Limitations and Recommended Next Steps", 1)
    add_number(doc, "Complete wheels-raised verification of forward, stop, left yaw, and right yaw commands.")
    add_number(doc, "Collect stationary ultrasonic samples at known distances (20, 30, 40, and 50 cm) and record variance and outliers.")
    add_number(doc, "Run supervised low-speed floor trials and measure stopping distance at several battery levels and payloads.")
    add_number(doc, "Calibrate turn duration and yaw rate against measured heading change; consider encoder or IMU feedback for closed-loop turning.")
    add_number(doc, "Connect the simplified camera detection supplier through turn_advisor only after the ultrasonic loop is physically validated.")
    add_number(doc, "Add integration logs that record timestamp, raw and filtered distance, state transition, motor command, and selected turn source.")
    add_number(doc, "Define a watchdog or heartbeat for the controller process and add a hardware-level emergency-stop mechanism.")
    add_number(doc, "Evaluate the four-channel line follower only when route-following becomes an explicit requirement.")

    add_heading(doc, "10. Conclusion", 1)
    add_paragraph(doc, "The vehicle code follows a sensible prototype strategy: start with a small, bounded local behavior; isolate hardware behind adapters; make every uncertain condition stop the vehicle; and add perception only through a constrained advisory interface. The current software tests support the control design, and the visual logs support the feasibility of near-real-time guidance. The next engineering milestone should focus on measured physical validation rather than adding more perception features.")

    add_heading(doc, "Appendix A. Code Reference Map", 1)
    add_table(doc,
        ["Repository path", "Report use"],
        [
            ("TurboPiIntegratedGuide/main.py", "Command-line entry point and simulation/hardware gate."),
            ("TurboPiIntegratedGuide/controller.py", "Finite-state controller, filtering, turn selection, and fail-safe stop."),
            ("TurboPiIntegratedGuide/hardware.py", "HiWonder SDK adapters and deterministic simulation objects."),
            ("TurboPiIntegratedGuide/config.py", "Validated vehicle tuning parameters."),
            ("TurboPiIntegratedGuide/visual_guidance.py", "Minimal camera-to-turn advisory boundary."),
            ("TurboPiIntegratedGuide/tests/test_controller.py", "Automated vehicle-control verification."),
            ("VisualGuideProject/detection_api.py", "Candidate selection and YOLO labeling boundary."),
            ("VisualGuideProject/guidance_api.py", "Region, direction, and closeness interpretation."),
            ("VisualGuideProject/risk_api.py", "Multi-frame approach-risk scoring."),
            ("VisualGuideProject/stability_api.py", "Detection and warning stabilization."),
            ("VisualGuideProject/distance_api.py", "Approximate monocular-distance formula."),
            ("VisualGuideProject/TEST_LOG.md", "Recorded live-camera performance evidence."),
        ], [4100, 5260])
    add_paragraph(doc, "All descriptions in this report were derived from the repository files listed above. No claim of completed physical-vehicle validation is made.", italic=True, color=GRAY)

    # Prevent automatic compression of embedded image and request field updates.
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    do_not_compress = OxmlElement("w:doNotAutoCompressPictures")
    settings.append(do_not_compress)

    doc.core_properties.title = "Vehicle Control Code and Strategy"
    doc.core_properties.subject = "TurboPi obstacle avoidance and visual-guidance implementation report"
    doc.core_properties.author = "Smart Vehicle Project"
    doc.core_properties.keywords = "TurboPi, obstacle avoidance, ultrasonic, computer vision, safety controller"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
